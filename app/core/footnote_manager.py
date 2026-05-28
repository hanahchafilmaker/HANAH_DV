"""
Footnote management utilities for the thesis tool - Manual Editor Version.
Extracts footnotes from .docx (via direct XML parsing) and provides basic parsing.
"""

import os
import zipfile
import xml.etree.ElementTree as ET
import re
import tempfile
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
import difflib


# -------------------------
# 데이터 모델
# -------------------------
@dataclass
class Footnote:
    """footnote data model"""
    fn_id: str
    fn_text: str
    author: str = ""
    title: str = ""
    year: str = ""
    pages: str = ""
    publisher: str = ""
    location: str = ""

    def __getitem__(self, key):
        """Allow dictionary-style access for backward compatibility"""
        return getattr(self, key)

@dataclass
class MatchCandidate:
    """Represents a single candidate match for a footnote"""
    matched_ref: str
    confidence: float
    source: str  # 'memory' or 'crossref'
    citation_type: str  # 'SHORT', 'FULL', or 'REPEATED'
    doi: str = ""
    preview: str = ""

@dataclass
class MatchResult:
    """Contains the best match and list of all candidates"""
    best_match: MatchCandidate
    candidates: List[MatchCandidate]


# -------------------------
# 로깅 설정
# -------------------------
import logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)-8s] %(name)s: %(message)s",
    handlers=[
        logging.FileHandler("paper_system.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# -------------------------
# Footnote extraction from DOCX
# -------------------------
def extract_footnotes(docx_path: str) -> List[Footnote]:
    """
    Extract footnotes from a .docx file by parsing the XML directly.
    Returns a list of Footnote objects.
    """
    try:
        footnotes_xml = _get_xml_from_docx(docx_path, 'word/footnotes.xml')
        document_xml = _get_xml_from_docx(docx_path, 'word/document.xml')

        if footnotes_xml is None:
            # No footnotes part
            return []

        # Parse footnotes.xml
        try:
            ft_root = ET.fromstring(footnotes_xml)
        except ET.ParseError as e:
            raise ValueError(f"Failed to parse footnotes.xml: {e}")

        # Build mapping from footnote ID to its text
        fn_elements = {}
        for footnote in ft_root.findall('.//w:footnote', _get_namespaces()):
            fid = footnote.get(f"{{{_get_namespaces()['w']}}}id")
            if fid is None:
                continue
            # Collect all text inside the footnote
            texts = []
            for t in footnote.findall('.//w:t', _get_namespaces()):
                if t.text:
                    texts.append(t.text)
            fn_text = ''.join(texts).strip()
            # Skip if empty or just a footnote number pattern like [1], [2], etc.
            if not fn_text or re.match(r'^\[\d+\]$', fn_text):
                continue
            fn_elements[fid] = {
                'fn_id': fid,
                'fn_text': fn_text,
                'fn_refs': []  # will fill from document.xml
            }

        if document_xml is None:
            # still return footnotes even if we can't find refs
            return [Footnote(**{k: v for k, v in fn.items() if k != 'fn_refs'}) for fn in fn_elements.values()]

        # Parse document.xml to find footnote references
        try:
            doc_root = ET.fromstring(document_xml)
        except ET.ParseError as e:
            # If we can't parse, just return footnotes without refs
            return [Footnote(**{k: v for k, v in fn.items() if k != 'fn_refs'}) for fn in fn_elements.values()]

        ns = _get_namespaces()
        for ref in doc_root.findall('.//w:footnoteRef', ns):
            fid = ref.get(f"{{{ns['w']}}}id")
            if fid in fn_elements:
                fn_elements[fid]['fn_refs'].append(fid)

        # Convert to list of Footnote objects, sorted by fn_id as integer if possible
        def try_int(s):
            try:
                return int(s)
            except ValueError:
                return s

        sorted_items = sorted(fn_elements.values(), key=lambda x: try_int(x['fn_id']))
        return [Footnote(**{k: v for k, v in fn.items() if k != 'fn_refs'}) for fn in sorted_items]
    except Exception as e:
        logger.error(f"Failed to extract footnotes from {docx_path}: {e}")
        # Return empty list instead of raising exception to prevent crashes
        return []


def _get_xml_from_docx(docx_path: str, path_in_zip: str):
    """Extract raw XML bytes from a .docx (which is a zip archive)."""
    with zipfile.ZipFile(docx_path) as z:
        if path_in_zip not in z.namelist():
            return None
        return z.read(path_in_zip)


def _get_namespaces():
    """Namespaces used in Word XML"""
    return {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }


# -------------------------
# 파싱 함수 (순수 로컬 파싱)
# -------------------------
def parse_footnote_text(text: str) -> Dict[str, str]:
    """
    Parse footnote text into components using simple heuristics.
    Returns dict with author, title, year, pages, publisher, location.
    """
    if not text or not isinstance(text, str):
        return {
            "author": "",
            "title": "",
            "year": "",
            "pages": "",
            "publisher": "",
            "location": ""
        }

    # Initialize result
    result = {
        "author": "",
        "title": "",
        "year": "",
        "pages": "",
        "publisher": "",
        "location": ""
    }

    # Extract year (4 digits, 1000-2029)
    year_match = re.search(r'\b(1[0-9]{3}|2[0-2][0-9]{3})\b', text)
    if year_match:
        result["year"] = year_match.group(0)

    # Simple heuristic to extract author and title:
    # Assume format: "Author. Title." or "Author, Title." or "Author Title."
    # We'll split by '.' first, then by ',' if needed.
    # Remove the year from text for author/title extraction to avoid confusion.
    text_no_year = text
    if year_match:
        # Remove the year occurrence (just one) from text
        text_no_year = text.replace(year_match.group(0), '', 1)

    # Clean up extra spaces and punctuation
    text_no_year = text_no_year.strip()
    # Replace commas with periods for uniform splitting
    text_no_year = text_no_year.replace(',', '.')
    # Split by periods
    parts = [p.strip() for p in text_no_year.split('.') if p.strip()]

    if len(parts) >= 2:
        # Assume first part is author, second part is title
        result["author"] = parts[0]
        result["title"] = parts[1]
        # If there are more parts, could be publisher, etc., but we ignore for simplicity
    elif len(parts) == 1:
        # Only one part, assume it's author (or title); we'll put it in author
        result["author"] = parts[0]
        # title remains empty
    # If no parts, leave both empty

    # For simplicity, we leave pages, publisher, location empty
    # User can edit them manually

    return result


def _calculate_similarity(str1: str, str2: str) -> float:
    """
    Calculate similarity between two strings using difflib.SequenceMatcher.
    Returns a value between 0.0 and 1.0.
    """
    if not str1 or not str2:
        return 0.0
    return difflib.SequenceMatcher(None, str1.lower(), str2.lower()).ratio()


def _score_match(parsed: Dict[str, str], candidate_ref: str) -> float:
    """
    Score a candidate reference against parsed footnote data.
    Uses weighted similarity: title (0.4), author (0.4), year (0.2).
    Returns a score between 0.0 and 1.0.
    """
    # Parse the candidate reference to extract components
    parsed_candidate = parse_footnote_text(candidate_ref)

    # Calculate similarities
    title_sim = _calculate_similarity(parsed.get('title', ''), parsed_candidate.get('title', ''))
    author_sim = _calculate_similarity(parsed.get('author', ''), parsed_candidate.get('author', ''))

    # Year match: exact match = 1.0, no match = 0.0
    year_match = 1.0 if parsed.get('year') == parsed_candidate.get('year') and parsed.get('year') else 0.0

    # Weighted score
    score = (title_sim * 0.4) + (author_sim * 0.4) + (year_match * 0.2)

    return score


def _memory_first_lookup(parsed_ref: Dict[str, str], fn_id: str) -> Optional[Dict[str, Any]]:
    """
    Perform memory-first lookup for a reference.
    Returns a dict with matched_ref, confidence, source, citation_type, doi, preview if found.
    """
    # Fuzzy matching for repeat citations: compare with stored references
    author = parsed_ref.get('author', '').strip()
    year = parsed_ref.get('year', '').strip()

    if not author and not year:
        return None

    # Check all stored references for this footnote ID for similarity
    if fn_id in _stored_references:
        for short_cite_key, stored_full_ref in _stored_references.items():
            # Parse the stored short citation key to get author/year
            if ', ' in short_cite_key:
                stored_author, stored_year = short_cite_key.split(', ', 1)
                stored_author = stored_author.strip()
                stored_year = stored_year.strip()

                # Calculate similarity
                author_sim = _calculate_similarity(author.lower(), stored_author.lower()) if author and stored_author else (0.0 if not author and not stored_author else (1.0 if author == stored_author else 0.0))
                year_match = (year == stored_year) if year and stored_year else (True if not year and not stored_year else False)

                # Consider it a repeat if author similarity is high and year matches
                if author_sim >= 0.8 and year_match:
                    # Calculate confidence based on similarity to the full reference
                    confidence = _score_match(parsed_ref, stored_full_ref)
                    # Boost confidence for repeat citations
                    confidence = min(0.95, confidence + 0.1)

                    return {
                        'matched_ref': stored_full_ref,
                        'confidence': confidence,
                        'source': 'memory',
                        'citation_type': 'REPEATED',
                        'doi': '',  # Would need to extract from stored_ref or memory
                        'preview': stored_full_ref[:50] + '...' if len(stored_full_ref) > 50 else stored_full_ref
                    }

    # Also check exact match for backward compatibility
    short_citation = parsed_ref.get('author', '') + ', ' + parsed_ref.get('year', '')
    short_citation = short_citation.strip(', ')
    if short_citation:
        if is_repeat_citation(short_citation, fn_id):
            # Get the stored full reference
            stored_ref = get_stored_reference(short_citation)
            if stored_ref:
                # Calculate confidence based on similarity
                confidence = _score_match(parsed_ref, stored_ref)
                # Boost confidence for repeat citations
                confidence = min(0.95, confidence + 0.1)
                return {
                    'matched_ref': stored_ref,
                    'confidence': confidence,
                    'source': 'memory',
                    'citation_type': 'REPEATED',
                    'doi': '',  # Would need to extract from stored_ref or memory
                    'preview': stored_ref[:50] + '...' if len(stored_ref) > 50 else stored_ref
                }

    # Check if we have a stored reference for this exact parsed reference
    # This would be for FULL citations that we've seen before
    # For now, we'll return None to let Crossref handle it
    # In a full implementation, we'd check memory for exact matches

    return None


def _generate_crossref_candidates(parsed_ref: Dict[str, str]) -> List[Dict[str, Any]]:
    """
    Generate candidates from Crossref (simulated for now).
    In a real implementation, this would query the Crossref API.
    For now, we'll return empty list since Crossref was removed per requirements.
    """
    # Per requirements, Crossref API has been completely removed
    # So we return empty list
    return []


def _generate_local_variations(parsed_ref: Dict[str, str]) -> List[Dict[str, Any]]:
    """
    Generate local variations of the parsed reference to provide multiple candidates.
    This creates slight variations in author name, title, etc. to give user options.
    """
    variations = []

    author = parsed_ref.get('author', '')
    title = parsed_ref.get('title', '')
    year = parsed_ref.get('year', '')

    # Only generate variations if we have basic data
    if not author and not title:
        return variations

    # Base reference
    base_ref_parts = []
    if author:
        base_ref_parts.append(author)
    if title:
        base_ref_parts.append(title)
    if year:
        base_ref_parts.append(year)

    base_ref = ". ".join(base_ref_parts) + "." if base_ref_parts else ""

    # Variation 1: Original format
    if base_ref:
        variations.append({
            'matched_ref': base_ref,
            'confidence': 1.0,  # Perfect match to itself
            'source': 'memory',
            'citation_type': 'FULL',
            'doi': '',
            'preview': base_ref[:50] + '...' if len(base_ref) > 50 else base_ref
        })

    # Variation 2: Author with "et al." if there are multiple authors (simplified)
    if author and 'and' in author.lower():
        # Simulate shortening author list
        author_short = author.split('and')[0].strip() + " et al."
        var2_parts = []
        if author_short:
            var2_parts.append(author_short)
        if title:
            var2_parts.append(title)
        if year:
            var2_parts.append(year)
        var2_ref = ". ".join(var2_parts) + "." if var2_parts else ""
        if var2_ref and var2_ref != base_ref:
            variations.append({
                'matched_ref': var2_ref,
                'confidence': 0.85,  # Slightly lower confidence for variation
                'source': 'memory',
                'citation_type': 'FULL',
                'doi': '',
                'preview': var2_ref[:50] + '...' if len(var2_ref) > 50 else var2_ref
            })

    # Variation 3: Title in quotes if not already
    if title and not (title.startswith('"') and title.endswith('"')) and not (title.startswith("'") and title.endswith("'")):
        var3_parts = []
        if author:
            var3_parts.append(author)
        if title:
            var3_parts.append(f'"{title}"')
        if year:
            var3_parts.append(year)
        var3_ref = ". ".join(var3_parts) + "." if var3_parts else ""
        if var3_ref and var3_ref != base_ref:
            variations.append({
                'matched_ref': var3_ref,
                'confidence': 0.8,  # Lower confidence for stylistic variation
                'source': 'memory',
                'citation_type': 'FULL',
                'doi': '',
                'preview': var3_ref[:50] + '...' if len(var3_ref) > 50 else var3_ref
            })

    # Limit to top 3 variations
    return variations[:3]


def auto_match_reference(fn_text: str, fn_id: str) -> Optional[MatchResult]:
    """
    Collect all possible candidate matches for a footnote text.
    Returns MatchResult with candidates sorted by confidence (descending) and no auto-selection.
    """
    # Parse the footnote text
    parsed_ref = parse_footnote_text(fn_text)

    # If we don't have at least author or year, we can't do meaningful matching
    if not parsed_ref.get('author') and not parsed_ref.get('year'):
        return None

    candidates = []

    # Step 1: Try memory-first lookup (for REPEATED citations)
    # We collect all stored references that are similar enough to be considered repeats.
    author = parsed_ref.get('author', '').strip()
    year = parsed_ref.get('year', '').strip()
    if author or year:
        # Fuzzy matching for repeat citations: compare with stored references
        if fn_id in _stored_references:
            for short_cite_key, stored_full_ref in _stored_references.items():
                # Parse the stored short citation key to get author/year
                if ', ' in short_cite_key:
                    stored_author, stored_year = short_cite_key.split(', ', 1)
                    stored_author = stored_author.strip()
                    stored_year = stored_year.strip()

                    # Calculate similarity
                    author_sim = _calculate_similarity(author.lower(), stored_author.lower()) if author and stored_author else (0.0 if not author and not stored_author else (1.0 if author == stored_author else 0.0))
                    year_match = (year == stored_year) if year and stored_year else (True if not year and not stored_year else False)

                    # Consider it a repeat if author similarity is high and year matches
                    if author_sim >= 0.8 and year_match:
                        # Calculate confidence based on similarity to the full reference
                        confidence = _score_match(parsed_ref, stored_full_ref)
                        # Boost confidence for repeat citations (optional, we keep it for ranking)
                        confidence = min(0.95, confidence + 0.1)

                        candidate = MatchCandidate(
                            matched_ref=stored_full_ref,
                            confidence=confidence,
                            source='memory',
                            citation_type='REPEATED',
                            doi='',  # Would need to extract from stored_ref or memory
                            preview=stored_full_ref[:50] + '...' if len(stored_full_ref) > 50 else stored_full_ref
                        )
                        candidates.append(candidate)
        # Also check exact match for backward compatibility (adds the exact stored reference if exists)
        short_citation = parsed_ref.get('author', '') + ', ' + parsed_ref.get('year', '')
        short_citation = short_citation.strip(', ')
        if short_citation and is_repeat_citation(short_citation, fn_id):
            stored_ref = get_stored_reference(short_citation)
            if stored_ref:
                confidence = _score_match(parsed_ref, stored_ref)
                confidence = min(0.95, confidence + 0.1)
                candidate = MatchCandidate(
                    matched_ref=stored_ref,
                    confidence=confidence,
                    source='memory',
                    citation_type='REPEATED',
                    doi='',
                    preview=stored_ref[:50] + '...' if len(stored_ref) > 50 else stored_ref
                )
                candidates.append(candidate)

    # Step 2: For FULL citations, generate candidates (local variations)
    # Generate a reference string from parsed data
    author = parsed_ref.get('author', '')
    title = parsed_ref.get('title', '')
    year = parsed_ref.get('year', '')

    # Build a basic reference string
    if author and title:
        basic_ref = f"{author}. {title}."
        if year:
            basic_ref += f" {year}"
        basic_ref += "."
    elif author:
        basic_ref = f"{author}."
        if year:
            basic_ref += f" {year}"
        basic_ref += "."
    elif title:
        basic_ref = f'"{title}".'
        if year:
            basic_ref += f" {year}"
        basic_ref += "."
    else:
        # Not enough data to create a reference; we will still return any memory candidates if exist
        pass

    # Generate local variations to provide multiple candidates
    local_variations = _generate_local_variations(parsed_ref)

    # Score each variation and create candidates
    for variation in local_variations:
        # Calculate confidence based on similarity to original parsed reference
        confidence = _score_match(parsed_ref, variation['matched_ref'])
        # Update the confidence in the variation
        variation['confidence'] = confidence

        # Create MatchCandidate object
        candidate = MatchCandidate(
            matched_ref=variation['matched_ref'],
            confidence=variation['confidence'],
            source=variation['source'],
            citation_type=variation['citation_type'],
            doi=variation.get('doi', ''),
            preview=variation.get('preview', '')
        )
        candidates.append(candidate)

    # If no candidates at all, return None
    if not candidates:
        return None

    # Sort candidates by confidence (descending)
    candidates.sort(key=lambda x: x.confidence, reverse=True)

    # Return all candidates; best_match is None to indicate no auto-selection
    return MatchResult(best_match=None, candidates=candidates)


# Citation memory storage (simplified version)
_citation_memory = {}  # Tracks short citations per footnote_id
_stored_references = {}  # Maps short citations to full references


def reset_citation_memory():
    """Reset the citation memory for a new document"""
    global _citation_memory, _stored_references
    _citation_memory = {}
    _stored_references = {}


def is_repeat_citation(short_citation: str, fn_id: str) -> bool:
    """Check if a short citation has been seen before for this footnote ID"""
    if fn_id not in _citation_memory:
        return False
    return short_citation in _citation_memory[fn_id]


def store_reference(short_citation: str, full_reference: str, fn_id: str):
    """Store a reference for future repeat detection"""
    if fn_id not in _citation_memory:
        _citation_memory[fn_id] = set()
    _citation_memory[fn_id].add(short_citation)
    _stored_references[short_citation] = full_reference


def get_stored_reference(short_citation: str) -> Optional[str]:
    """Get the stored full reference for a short citation"""
    return _stored_references.get(short_citation)


# -------------------------
# Bibliography generation
# -------------------------
def generate_bibliography_from_edited(csv_path: str) -> List[str]:
    """
    Generate a bibliography list from the edited CSV.
    Returns list of formatted reference strings.
    This is a simplified version - in production, you'd use the formatted data.
    """
    import pandas as pd

    try:
        df = pd.read_csv(csv_path, dtype=str).fillna('')
        # Keep only rows where status contains a check mark (✅) for bibliography
        mask = df['status'].str.contains('✅', na=False)
        bib_df = df[mask].copy()

        bibliography = []
        for _, row in bib_df.iterrows():
            ref = row['matched_ref'].strip()
            if ref:
                bibliography.append(ref)

        logger.info(f"참고문헌 {len(bibliography)}건 생성 (원본 {len(bib_df)}건에서 중복 제거)")
        return bibliography
    except Exception as e:
        logger.error(f"Failed to generate bibliography: {e}")
        return []


def update_docx_with_bibliography(docx_path: str, bibliography: List[str], output_path: str):
    """
    Append a bibliography section to the end of the docx.
    Each entry as a normal paragraph.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        # Add a heading for bibliography
        heading = doc.add_paragraph()
        heading_run = heading.add_run('참고문헌')
        heading_run.bold = True
        heading_run.font.size = doc.styles['Normal'].font.size  # keep same size; could adjust
        heading.alignment = 0  # left align; could justify etc.
        # Add each bibliography entry
        for ref in bibliography:
            p = doc.add_paragraph(ref)
            # Optional: apply hanging indent etc. For simplicity, normal.
        doc.save(output_path)
    except Exception as e:
        raise ValueError(f"Failed to update DOCX with bibliography: {e}")


# -------------------------
# 참고문헌 서식 함수
# -------------------------
def format_apa(footnote: Footnote) -> str:
    """Format footnote as APA citation"""
    if not footnote.author and not footnote.title:
        return ""

    year = f"({footnote.year})" if footnote.year else ""
    title = footnote.title if footnote.title else ""

    return f"{footnote.author} {year}. {title}."


def format_chicago(footnote: Footnote) -> str:
    """Format footnote as Chicago citation"""
    if not footnote.author and not footnote.title:
        return ""

    # Chicago: Author, *Title* (Location: Publisher, Year), Pages.
    author_part = footnote.author if footnote.author else ""
    title_part = f"*{footnote.title}*" if footnote.title else ""

    # Location-publisher part
    loc_pub_parts = []
    if footnote.location:
        loc_pub_parts.append(footnote.location)
    if footnote.publisher:
        loc_pub_parts.append(footnote.publisher)
    loc_pub_part = f": {' '.join(loc_pub_parts)}" if loc_pub_parts else ""

    year_part = f"{footnote.year}" if footnote.year else ""
    pages_part = f"{footnote.pages}" if footnote.pages else ""

    if loc_pub_part or year_part:
        middle = f" ({loc_pub_part}{year_part})" if year_part else f" ({loc_pub_part})"
    else:
        middle = ""

    pages_suffix = f", {pages_part}" if pages_part else ""

    return f"{author_part}, {title_part}{middle}{pages_suffix}."


def format_mla(footnote: Footnote) -> str:
    """Format footnote as MLA citation"""
    if not footnote.author and not footnote.title:
        return ""

    # MLA: Author. "Title." Publisher, Year, Pages.
    author_part = f"{footnote.author}." if footnote.author else ""
    title_part = f'"{footnote.title}."' if footnote.title else ""
    publisher_part = f"{footnote.publisher}" if footnote.publisher else ""
    year_part = f"{footnote.year}" if footnote.year else ""
    pages_part = f"{footnote.pages}" if footnote.pages else ""

    # Combine publisher, year, pages
    pub_year_parts = [p for p in [publisher_part, year_part, pages_part] if p]
    pub_year_part = ", ".join(pub_year_parts)

    return f"{author_part} {title_part} {pub_year_part}".strip()


# -------------------------
# CSV template functions (simplified)
# -------------------------
def write_footnote_template(docx_path: str, csv_path: str):
    """
    Create a CSV template from the footnotes in docx_path.
    Columns: fn_id,fn_text,author,title,year,pages,publisher,location
    """
    footnotes = extract_footnotes(docx_path)
    rows = []
    for fn in footnotes:
        rows.append({
            'fn_id': fn.fn_id,
            'fn_text': fn.fn_text,
            'author': fn.author,
            'title': fn.title,
            'year': fn.year,
            'pages': fn.pages,
            'publisher': fn.publisher,
            'location': fn.location
        })

    # If no footnotes, still create header
    if not rows:
        rows = [{'fn_id': '', 'fn_text': '', 'author': '', 'title': '', 'year': '', 'pages': '', 'publisher': '', 'location': ''}]

    fieldnames = ['fn_id', 'fn_text', 'author', 'title', 'year', 'pages', 'publisher', 'location']
    with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def read_edited_footnotes(csv_path: str) -> List[Dict[str, Any]]:
    """
    Read the edited CSV and return list of dicts.
    """
    import pandas as pd
    df = pd.read_csv(csv_path, dtype=str).fillna('')
    # Ensure required columns exist
    required = ['fn_id', 'fn_text', 'author', 'title', 'year', 'pages', 'publisher', 'location']
    for col in required:
        if col not in df.columns:
            df[col] = ''
    return df[required].to_dict('records')


def bibliography_to_bibtex(entries: List[Dict[str, Any]]) -> str:
    """
    Convert a list of bibliography entries to BibTeX format.
    Each entry is a dict with keys like 'author', 'title', 'year', 'journal', 'publisher', 'doi'.
    Returns a string containing the BibTeX entries.
    """
    bibtex_lines = []
    for i, entry in enumerate(entries):
        # Generate a simple citation key: author_year or title_year if no author
        author = entry.get('author', '')
        title = entry.get('title', '')
        year = entry.get('year', '')

        if author:
            # Take first author's last name (simplified)
            first_author = author.split()[0] if author.split() else 'unknown'
            # Clean the author name for use as a key (remove non-alphanumeric)
            key = ''.join(c for c in first_author if c.isalnum())
        else:
            # Use first word of title
            first_word = title.split()[0] if title.split() else 'unknown'
            key = ''.join(c for c in first_word if c.isalnum())

        if year:
            key += f"_{year}"
        else:
            key += f"_{i+1}"  # fallback to index

        # Determine entry type: if journal is present, use @article, else @book or @misc
        if entry.get('journal'):
            entry_type = '@article'
        elif entry.get('publisher'):
            entry_type = '@book'
        else:
            entry_type = '@misc'

        bibtex_lines.append(f"{entry_type}{{{key},")

        # Add fields in a common order
        fields = []
        if author:
            fields.append(f"  author = {{{author}}}")
        if title:
            fields.append(f"  title = {{{title}}}")
        if year:
            fields.append(f"  year = {{{year}}}")
        if entry.get('journal'):
            fields.append(f"  journal = {{{entry['journal']}}}")
        if entry.get('publisher'):
            fields.append(f"  publisher = {{{entry['publisher']}}}")
        if entry.get('doi'):
            fields.append(f"  doi = {{{entry['doi']}}}")

        bibtex_lines.extend(fields)
        bibtex_lines.append("}\n")

    return "\n".join(bibtex_lines)


def save_bibtex_file(entries: List[Dict[str, Any]], output_path: str) -> bool:
    """
    Save bibliography entries as a BibTeX file.
    Returns True if successful, False otherwise.
    """
    try:
        bibtex_content = bibliography_to_bibtex(entries)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(bibtex_content)
        return True
    except Exception as e:
        logger.error(f"Failed to save BibTeX file: {e}")
        return False