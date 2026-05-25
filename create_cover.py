from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================
# 🔧 입력값 (여기만 수정)
# =========================
TITLE = "논문 제목"
SUBTITLE = "부제목 (없으면 빈 문자열)"
AUTHOR = "차하나"
MAJOR = "영화이론전공"
PROF = "지도교수 ○○○"
YEAR = "2026"

OUTPUT = "표지.docx"


# =========================
# 함수
# =========================
def add_paragraph(doc, text, size, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

    p.paragraph_format.space_after = Pt(space_after)
    return p


# =========================
# 표지 생성
# =========================
def create_cover():

    doc = Document()
    section = doc.sections[0]

    # ✅ 논문 규격 여백
    section.top_margin = Cm(5.5)
    section.bottom_margin = Cm(5.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(4)

    # =========================
    # 구성
    # =========================
    add_paragraph(doc, "학 위 논 문", 16, True, 50)

    add_paragraph(doc, TITLE, 20, True, 30)

    if SUBTITLE:
        add_paragraph(doc, SUBTITLE, 14, False, 40)

    add_paragraph(doc, PROF, 14, False, 40)

    add_paragraph(doc, AUTHOR, 16, False, 20)
    add_paragraph(doc, MAJOR, 14, False, 30)

    add_paragraph(doc, YEAR, 14, False, 0)

    # 저장
    doc.save(OUTPUT)
    print("✅ 표지 생성 완료:", OUTPUT)


# =========================
# 실행
# =========================
if __name__ == "__main__":
    create_cover()