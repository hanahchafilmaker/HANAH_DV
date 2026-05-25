from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================
# 🔧 입력값 (여기 수정)
# =========================
TITLE = "논문 제목"
AUTHOR = "차하나"
DEGREE = "석사학위 논문"
PROF = "지도교수 ○○○"

COMMITTEE = [
    "위원장 ○○○",
    "위원 ○○○",
    "위원 ○○○",
]

DATE_SUBMIT = "2026년 6월"
DATE_APPROVE = "2026년 7월"

OUTPUT = "인준지.docx"


# =========================
# 함수
# =========================
def add_paragraph(doc, text, size, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

    p.paragraph_format.space_after = Pt(space_after)
    return p


# =========================
# 인준지 생성
# =========================
def create_approval():

    doc = Document()
    section = doc.sections[0]

    # ✅ 논문 규격 여백
    section.top_margin = Cm(5.5)
    section.bottom_margin = Cm(5.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(4)

    # =========================
    # 구성
    # =========================
    add_paragraph(doc, DEGREE, 16, True, 40)

    add_paragraph(doc, TITLE, 18, True, 30)

    add_paragraph(doc, AUTHOR, 14, False, 20)
    add_paragraph(doc, PROF, 14, False, 30)

    # 제출 / 인준 날짜
    add_paragraph(doc, f"제출일: {DATE_SUBMIT}", 12, False, 10)
    add_paragraph(doc, f"인준일: {DATE_APPROVE}", 12, False, 30)

    # 위원
    for c in COMMITTEE:
        add_paragraph(doc, f"{c} (인)", 12, False, 12)

    add_paragraph(doc, "동국대학교 대학원", 14, True, 0)

    # 저장
    doc.save(OUTPUT)
    print("✅ 인준지 생성 완료:", OUTPUT)


# =========================
# 실행
# =========================
if __name__ == "__main__":
    create_approval()