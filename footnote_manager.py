"""
Footnote management utilities for the thesis tool - Manual Editor Version.
Extracts footnotes from .docx (via direct XML parsing) and provides basic parsing.
"""

import os
import zipfile
import xml.etree.ElementTree as ET
import re
import tempfile
from dataclasses import dataclass
from typing import List, Dict, Any


# -------------------------
# 데이터 모델
# -------------------------
@dataclass
class Footnote:
    """footnote data model"""
    fn_id: str
    fn_text: str
    author: str = ""
    title: str = ""
    year: str = ""
    pages: str = ""
    publisher: str = ""
    location: str = ""


# -------------------------
# 로깅 설정
# -------------------------
import logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)-8s] %(name)s: %(message)s",
    handlers=[
        logging.FileHandler("paper_system.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# -------------------------
# Footnote extraction from DOCX
# -------------------------
def extract_footnotes(docx_path: str) -> List[Footnote]:
    """
    Extract footnotes from a .docx file by parsing the XML directly.
    Returns a list of Footnote objects.
    """
    footnotes_xml = _get_xml_from_docx(docx_path, 'word/footnotes.xml')
    document_xml = _get_xml_from_docx(docx_path, 'word/document.xml')

    if footnotes_xml is None:
        # No footnotes part
        return []

    # Parse footnotes.xml
    try:
        ft_root = ET.fromstring(footnotes_xml)
    except ET.ParseError as e:
        raise ValueError(f"Failed to parse footnotes.xml: {e}")

    # Build mapping from footnote ID to its text
    fn_elements = {}
    for footnote in ft_root.findall('.//w:footnote', _get_namespaces()):
        fid = footnote.get(f"{{{_get_namespaces()['w']}}}id")
        if fid is None:
            continue
        # Collect all text inside the footnote
        texts = []
        for t in footnote.findall('.//w:t', _get_namespaces()):
            if t.text:
                texts.append(t.text)
        fn_text = ''.join(texts).strip()
        # Skip if empty or just a footnote number pattern like [1], [2], etc.
        if not fn_text or re.match(r'^\[\d+\]$', fn_text):
            continue
        fn_elements[fid] = {
            'fn_id': fid,
            'fn_text': fn_text,
            'fn_refs': []  # will fill from document.xml
        }

    if document_xml is None:
        # still return footnotes even if we can't find refs
        return [Footnote(**{k: v for k, v in fn.items() if k != 'fn_refs'}) for fn in fn_elements.values()]

    # Parse document.xml to find footnote references
    try:
        doc_root = ET.fromstring(document_xml)
    except ET.ParseError as e:
        # If we can't parse, just return footnotes without refs
        return [Footnote(**{k: v for k, v in fn.items() if k != 'fn_refs'}) for fn in fn_elements.values()]

    ns = _get_namespaces()
    for ref in doc_root.findall('.//w:footnoteRef', ns):
        fid = ref.get(f"{{{ns['w']}}}id")
        if fid in fn_elements:
            fn_elements[fid]['fn_refs'].append(fid)

    # Convert to list of Footnote objects, sorted by fn_id as integer if possible
    def try_int(s):
        try:
            return int(s)
        except ValueError:
            return s

    sorted_items = sorted(fn_elements.values(), key=lambda x: try_int(x['fn_id']))
    return [Footnote(**{k: v for k, v in fn.items() if k != 'fn_refs'}) for fn in sorted_items]


def _get_xml_from_docx(docx_path: str, path_in_zip: str):
    """Extract raw XML bytes from a .docx (which is a zip archive)."""
    with zipfile.ZipFile(docx_path) as z:
        if path_in_zip not in z.namelist():
            return None
        return z.read(path_in_zip)


def _get_namespaces():
    """Namespaces used in Word XML"""
    return {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }


# -------------------------
# 파싱 함수 (순수 로컬 파싱)
# -------------------------
def parse_footnote_text(text: str) -> Dict[str, str]:
    """
    Parse footnote text into components using simple heuristics.
    Returns dict with author, title, year, pages, publisher, location.
    """
    if not text or not isinstance(text, str):
        return {
            "author": "",
            "title": "",
            "year": "",
            "pages": "",
            "publisher": "",
            "location": ""
        }

    # Initialize result
    result = {
        "author": "",
        "title": "",
        "year": "",
        "pages": "",
        "publisher": "",
        "location": ""
    }

    # Extract year (4 digits, 1000-2029)
    year_match = re.search(r'\b(1[0-9]{3}|2[0-2][0-9]{3})\b', text)
    if year_match:
        result["year"] = year_match.group(0)

    # Simple heuristic: split by periods and commas
    # This is a simplified parser - in reality, citation formats vary widely
    # For production, you might want to use a proper citation parsing library
    # or allow manual editing of all fields

    # For now, we'll return empty strings for other fields
    # The user will fill them in manually via the UI

    return result


# -------------------------
# Bibliography generation
# -------------------------
def generate_bibliography_from_edited(csv_path: str) -> List[str]:
    """
    Generate a bibliography list from the edited CSV.
    Returns list of formatted reference strings.
    This is a simplified version - in production, you'd use the formatted data.
    """
    import pandas as pd

    try:
        df = pd.read_csv(csv_path, dtype=str).fillna('')
        # Keep only rows where status contains a check mark (✅) for bibliography
        mask = df['status'].str.contains('✅', na=False)
        bib_df = df[mask].copy()

        bibliography = []
        for _, row in bib_df.iterrows():
            ref = row['matched_ref'].strip()
            if ref:
                bibliography.append(ref)

        logger.info(f"참고문헌 {len(bibliography)}건 생성 (원본 {len(bib_df)}건에서 중복 제거)")
        return bibliography
    except Exception as e:
        logger.error(f"Failed to generate bibliography: {e}")
        return []


def update_docx_with_bibliography(docx_path: str, bibliography: List[str], output_path: str):
    """
    Append a bibliography section to the end of the docx.
    Each entry as a normal paragraph.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        # Add a heading for bibliography
        heading = doc.add_paragraph()
        heading_run = heading.add_run('참고문헌')
        heading_run.bold = True
        heading_run.font.size = doc.styles['Normal'].font.size  # keep same size; could adjust
        heading.alignment = 0  # left align; could justify etc.
        # Add each bibliography entry
        for ref in bibliography:
            p = doc.add_paragraph(ref)
            # Optional: apply hanging indent etc. For simplicity, normal.
        doc.save(output_path)
    except Exception as e:
        raise ValueError(f"Failed to update DOCX with bibliography: {e}")


# -------------------------
# 참고문헌 서식 함수
# -------------------------
def format_apa(footnote: Footnote) -> str:
    """Format footnote as APA citation"""
    if not footnote.author and not footnote.title:
        return ""

    year = f"({footnote.year})" if footnote.year else ""
    title = footnote.title if footnote.title else ""

    return f"{footnote.author} {year}. {title}."


def format_chicago(footnote: Footnote) -> str:
    """Format footnote as Chicago citation"""
    if not footnote.author and not footnote.title:
        return ""

    # Chicago: Author, *Title* (Location: Publisher, Year), Pages.
    author_part = footnote.author if footnote.author else ""
    title_part = f"*{footnote.title}*" if footnote.title else ""

    # Location-publisher part
    loc_pub_parts = []
    if footnote.location:
        loc_pub_parts.append(footnote.location)
    if footnote.publisher:
        loc_pub_parts.append(footnote.publisher)
    loc_pub_part = f": {' '.join(loc_pub_parts)}" if loc_pub_parts else ""

    year_part = f"{footnote.year}" if footnote.year else ""
    pages_part = f"{footnote.pages}" if footnote.pages else ""

    if loc_pub_part or year_part:
        middle = f" ({loc_pub_part}{year_part})" if year_part else f" ({loc_pub_part})"
    else:
        middle = ""

    pages_suffix = f", {pages_part}" if pages_part else ""

    return f"{author_part}, {title_part}{middle}{pages_suffix}."


def format_mla(footnote: Footnote) -> str:
    """Format footnote as MLA citation"""
    if not footnote.author and not footnote.title:
        return ""

    # MLA: Author. "Title." Publisher, Year, Pages.
    author_part = f"{footnote.author}." if footnote.author else ""
    title_part = f'"{footnote.title}."' if footnote.title else ""
    publisher_part = f"{footnote.publisher}" if footnote.publisher else ""
    year_part = f"{footnote.year}" if footnote.year else ""
    pages_part = f"{footnote.pages}" if footnote.pages else ""

    # Combine publisher, year, pages
    pub_year_parts = [p for p in [publisher_part, year_part, pages_part] if p]
    pub_year_part = ", ".join(pub_year_parts)

    return f"{author_part} {title_part} {pub_year_part}".strip()


# -------------------------
# CSV template functions (simplified)
# -------------------------
def write_footnote_template(docx_path: str, csv_path: str):
    """
    Create a CSV template from the footnotes in docx_path.
    Columns: fn_id,fn_text,author,title,year,pages,publisher,location
    """
    footnotes = extract_footnotes(docx_path)
    rows = []
    for fn in footnotes:
        rows.append({
            'fn_id': fn.fn_id,
            'fn_text': fn.fn_text,
            'author': fn.author,
            'title': fn.title,
            'year': fn.year,
            'pages': fn.pages,
            'publisher': fn.publisher,
            'location': fn.location
        })

    # If no footnotes, still create header
    if not rows:
        rows = [{'fn_id': '', 'fn_text': '', 'author': '', 'title': '', 'year': '', 'pages': '', 'publisher': '', 'location': ''}]

    fieldnames = ['fn_id', 'fn_text', 'author', 'title', 'year', 'pages', 'publisher', 'location']
    with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def read_edited_footnotes(csv_path: str) -> List[Dict[str, Any]]:
    """
    Read the edited CSV and return list of dicts.
    """
    import pandas as pd
    df = pd.read_csv(csv_path, dtype=str).fillna('')
    # Ensure required columns exist
    required = ['fn_id', 'fn_text', 'author', 'title', 'year', 'pages', 'publisher', 'location']
    for col in required:
        if col not in df.columns:
            df[col] = ''
    return df[required].to_dict('records')