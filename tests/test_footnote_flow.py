import os
import tempfile
import footnote_manager as fm
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_test_docx(path):
    """Create a simple .docx with two footnotes."""
    doc = Document()
    # Add some text with footnote references
    p = doc.add_paragraph('This is a test sentence.')
    # Add a footnote reference (low-level)
    # We'll use the document's part.footnotes API if available; else fallback.
    # Simpler: just add a paragraph and then manually add footnote via docx's API?
    # python-docx does not have a high-level add_footnote yet (as of 0.8.10).
    # We'll use the low-level method from the python-docx documentation.
    # However, for testing extraction, we can rely on the fact that our extract_footnotes
    # reads from the footnotes part; we need to insert a w:footnote element.
    # Let's use the method from https://github.com/python-openxml/python-docx/issues/99
    def add_footnote(document, text):
        # Access the footnotes part
        footnotes_part = document.part.footnotes
        if footnotes_part is None:
            footnotes_part = document.part.footnotes = document.part.create_footnotes_part()
        # Create footnote element
        footnote = OxmlElement('w:footnote')
        footnote.set(qn('w:id'), str(len(footnotes_part._footnotes) + 1))
        # Add paragraph with text
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p.append(r)
        footnote.append(p)
        footnotes_part._footnotes.append(footnote)
        # Now we need to reference this footnote from the main document
        # We'll add a footnote reference run after the paragraph we want.
        # For simplicity, we'll just return the footnote id.
        return footnote.get(qn('w:id'))

    # Add a paragraph and then insert a footnote reference after it.
    p1 = doc.add_paragraph('First sentence that needs a footnote.')
    # Add a footnote reference run
    run = p1.add_run()
    # Create footnote reference element
    footnote_ref = OxmlElement('w:footnoteRef')
    footnote_ref.set(qn('w:id'), '1')  # placeholder; we will set after creating footnote
    run._r.append(footnote_ref)
    # Now create the footnote
    footnote_id = add_footnote(doc, 'This is the first footnote text.')
    # Update the reference with the correct id
    footnote_ref.set(qn('w:id'), footnote_id)

    # Second footnote
    p2 = doc.add_paragraph('Second sentence with another footnote.')
    run2 = p2.add_run()
    footnote_ref2 = OxmlElement('w:footnoteRef')
    footnote_ref2.set(qn('w:id'), '2')
    run2._r.append(footnote_ref2)
    footnote_id2 = add_footnote(doc, 'This is the second footnote text, maybe a bit longer.')
    footnote_ref2.set(qn('w:id'), footnote_id2)

    doc.save(path)
    return path

def test_extraction_and_bibliography():
    print("Creating test DOCX...")
    test_docx = "test_footnotes.docx"
    if os.path.exists(test_docx):
        os.remove(test_docx)
    create_test_docx(test_docx)
    print(f"Created {test_docx}")

    print("\nExtracting footnotes...")
    footnotes = fm.extract_footnotes(test_docx)
    print(f"Found {len(footnotes)} footnotes:")
    for fn in footnotes:
        print(f"  ID {fn['fn_id']}: {fn['fn_text'][:50]}...")

    print("\nGenerating CSV template...")
    csv_path = "test_footnotes_template.csv"
    fm.write_footnote_template(test_docx, csv_path)
    print(f"CSV written to {csv_path}")
    # Show CSV content
    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        print("CSV content:")
        print(f.read())

    print("\nSimulating user editing: fill matched_ref and set status...")
    # Read the CSV, modify rows to simulate user input
    import pandas as pd
    df = pd.read_csv(csv_path, dtype=str).fillna('')
    # For each row, set matched_ref to a formatted version and status to ✅ 매칭
    for i, row in df.iterrows():
        # Simple transformation: just wrap in quotes and add a period.
        df.at[i, 'matched_ref'] = f'"{row["fn_text"]}".'
        df.at[i, 'status'] = '✅ 매칭'
    df.to_csv(csv_path, index=False, encoding='utf-8-sig')
    print("Updated CSV (simulated user edits):")
    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        print(f.read())

    print("\nGenerating bibliography from edited CSV...")
    bibliography = fm.generate_bibliography_from_edited(csv_path)
    print(f"Generated {len(bibliography)} bibliography entries:")
    for i, ref in enumerate(bibliography, start=1):
        print(f"  {i}. {ref}")

    print("\nTesting update_docx_with_bibliography...")
    output_docx = "test_output_with_bibliography.docx"
    fm.update_docx_with_bibliography(test_docx, bibliography, output_docx)
    print(f"Output docx saved as {output_docx}")
    # Verify that the bibliography section exists by reading the docx and checking for the heading
    doc_out = Document(output_docx)
    # Look for a paragraph with text '참고문헌'
    found = False
    for para in doc_out.paragraphs:
        if para.text.strip() == '참고문헌':
            found = True
            print("Found bibliography heading.")
            # Print next few paragraphs as bibliography entries
            idx = doc_out.paragraphs.index(para)
            for j in range(1, min(4, len(doc_out.paragraphs) - idx)):
                next_para = doc_out.paragraphs[idx + j]
                if next_para.text.strip():
                    print(f"  Bibliography entry {j}: {next_para.text[:80]}...")
            break
    if not found:
        print("Warning: bibliography heading not found in output docx.")

    print("\nCleaning up test files...")
    for f in [test_docx, csv_path, output_docx]:
        if os.path.exists(f):
            try:
                os.remove(f)
                print(f"Removed {f}")
            except:
                pass
    print("Test completed.")

if __name__ == "__main__":
    test_extraction_and_bibliography()